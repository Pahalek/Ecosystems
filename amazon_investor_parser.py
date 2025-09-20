"""
Amazon Investor Reports Parser
Extracts key business metrics from Amazon's official investor report files (PDF, Excel, Word).
Focuses on segment data, operational metrics, and strategic information.
"""
import os
import re
import pandas as pd
import numpy as np
from datetime import datetime
import json
from pathlib import Path

# PDF processing
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF

# Excel processing
import openpyxl
from openpyxl import load_workbook

# Word processing
from docx import Document

# Tabular data extraction
try:
    import tabula
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False

class AmazonInvestorReportsParser:
    def __init__(self, reports_dir="investor_reports"):
        self.reports_dir = Path(reports_dir)
        self.extracted_data = {}
        self.parsed_files = []
        
        # Create subdirectories for organization
        (self.reports_dir / "10k_annual").mkdir(parents=True, exist_ok=True)
        (self.reports_dir / "10q_quarterly").mkdir(parents=True, exist_ok=True)
        (self.reports_dir / "earnings_releases").mkdir(parents=True, exist_ok=True)
        (self.reports_dir / "shareholder_letters").mkdir(parents=True, exist_ok=True)
        (self.reports_dir / "parsed_data").mkdir(parents=True, exist_ok=True)
        
        # Key metrics to extract
        self.target_metrics = {
            'aws_data': {
                'patterns': [
                    r'AWS.*revenue.*\$?([\d,\.]+).*billion',
                    r'Amazon Web Services.*\$?([\d,\.]+)',
                    r'AWS.*operating income.*\$?([\d,\.]+)',
                    r'AWS.*growth.*(\d+\.?\d*)%'
                ],
                'tables': ['AWS', 'Web Services', 'Segment']
            },
            'geographic_data': {
                'patterns': [
                    r'North America.*revenue.*\$?([\d,\.]+)',
                    r'International.*revenue.*\$?([\d,\.]+)',
                    r'North America.*operating.*\$?([\d,\.]+)',
                    r'International.*operating.*\$?([\d,\.]+)'
                ],
                'tables': ['Geographic', 'North America', 'International', 'Segment']
            },
            'advertising_data': {
                'patterns': [
                    r'advertising.*revenue.*\$?([\d,\.]+)',
                    r'advertising services.*\$?([\d,\.]+)',
                    r'advertising.*growth.*(\d+\.?\d*)%'
                ],
                'tables': ['Advertising', 'Services']
            },
            'prime_data': {
                'patterns': [
                    r'Prime.*member.*(\d+\.?\d*).*million',
                    r'Prime.*subscriber.*(\d+\.?\d*)',
                    r'(\d+).*million.*Prime.*member'
                ],
                'tables': ['Prime', 'Membership', 'Subscription']
            },
            'operational_metrics': {
                'patterns': [
                    r'fulfillment.*center.*(\d+)',
                    r'(\d+).*fulfillment.*center',
                    r'delivery.*station.*(\d+)',
                    r'third.*party.*seller.*(\d+\.?\d*)%',
                    r'marketplace.*(\d+\.?\d*)%'
                ],
                'tables': ['Fulfillment', 'Operations', 'Logistics']
            }
        }
    
    def parse_pdf_file(self, file_path):
        """Extract text and tables from PDF files."""
        print(f"📄 Parsing PDF: {file_path}")
        
        extracted = {
            'file_name': file_path.name,
            'file_type': 'PDF',
            'extracted_text': '',
            'tables': [],
            'metrics': {}
        }
        
        try:
            # Method 1: PyMuPDF for text extraction
            doc = fitz.open(file_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                full_text += page.get_text()
            
            extracted['extracted_text'] = full_text
            doc.close()
            
            # Method 2: pdfplumber for table extraction
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    # Extract tables
                    tables = page.extract_tables()
                    for j, table in enumerate(tables):
                        if table and len(table) > 1:  # Valid table
                            df = pd.DataFrame(table[1:], columns=table[0])
                            # Check if table contains relevant keywords
                            table_text = str(df).lower()
                            for metric_type, metric_info in self.target_metrics.items():
                                for keyword in metric_info['tables']:
                                    if keyword.lower() in table_text:
                                        extracted['tables'].append({
                                            'page': i+1,
                                            'table_index': j+1,
                                            'type': metric_type,
                                            'data': df.to_dict('records'),
                                            'keyword_matched': keyword
                                        })
                                        break
            
            # Extract specific metrics using patterns
            extracted['metrics'] = self.extract_metrics_from_text(full_text)
            
        except Exception as e:
            print(f"❌ Error parsing PDF {file_path}: {e}")
            
        return extracted
    
    def parse_excel_file(self, file_path):
        """Extract data from Excel files."""
        print(f"📊 Parsing Excel: {file_path}")
        
        extracted = {
            'file_name': file_path.name,
            'file_type': 'Excel',
            'sheets': {},
            'metrics': {}
        }
        
        try:
            workbook = load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Convert sheet to DataFrame
                data = []
                for row in sheet.iter_rows(values_only=True):
                    data.append(row)
                
                if data:
                    # Find header row (first non-empty row)
                    header_row = 0
                    for i, row in enumerate(data):
                        if any(cell is not None for cell in row):
                            header_row = i
                            break
                    
                    # Create DataFrame
                    df = pd.DataFrame(data[header_row+1:], columns=data[header_row])
                    df = df.dropna(how='all')  # Remove empty rows
                    
                    extracted['sheets'][sheet_name] = {
                        'data': df.to_dict('records'),
                        'shape': df.shape
                    }
                    
                    # Check for relevant metrics in this sheet
                    sheet_text = str(df).lower()
                    for metric_type, metric_info in self.target_metrics.items():
                        for keyword in metric_info['tables']:
                            if keyword.lower() in sheet_text:
                                extracted['metrics'][f"{sheet_name}_{metric_type}"] = {
                                    'sheet': sheet_name,
                                    'type': metric_type,
                                    'keyword_matched': keyword,
                                    'data_preview': df.head().to_dict('records')
                                }
                                break
            
        except Exception as e:
            print(f"❌ Error parsing Excel {file_path}: {e}")
            
        return extracted
    
    def parse_word_file(self, file_path):
        """Extract text from Word documents."""
        print(f"📝 Parsing Word: {file_path}")
        
        extracted = {
            'file_name': file_path.name,
            'file_type': 'Word',
            'extracted_text': '',
            'metrics': {}
        }
        
        try:
            doc = Document(file_path)
            full_text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + "\t"
                    full_text += row_text + "\n"
            
            extracted['extracted_text'] = full_text
            extracted['metrics'] = self.extract_metrics_from_text(full_text)
            
        except Exception as e:
            print(f"❌ Error parsing Word {file_path}: {e}")
            
        return extracted
    
    def extract_metrics_from_text(self, text):
        """Extract specific business metrics from text using regex patterns."""
        metrics = {}
        
        for metric_type, metric_info in self.target_metrics.items():
            metrics[metric_type] = []
            
            for pattern in metric_info['patterns']:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    metrics[metric_type].append({
                        'pattern': pattern,
                        'match': match.group(0),
                        'value': match.group(1) if match.groups() else None,
                        'context': text[max(0, match.start()-100):match.end()+100]
                    })
        
        return metrics
    
    def process_all_files(self):
        """Process all files in the investor_reports directory."""
        print("🔍 PROCESSING AMAZON INVESTOR REPORT FILES")
        print("=" * 60)
        
        if not self.reports_dir.exists():
            print(f"❌ Reports directory not found: {self.reports_dir}")
            return
        
        supported_extensions = {'.pdf', '.xlsx', '.xls', '.docx', '.doc'}
        processed_files = []
        
        # Process files in all subdirectories
        for file_path in self.reports_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    if file_path.suffix.lower() == '.pdf':
                        extracted = self.parse_pdf_file(file_path)
                    elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                        extracted = self.parse_excel_file(file_path)
                    elif file_path.suffix.lower() in ['.docx', '.doc']:
                        extracted = self.parse_word_file(file_path)
                    else:
                        continue
                    
                    self.extracted_data[file_path.name] = extracted
                    processed_files.append(file_path.name)
                    
                except Exception as e:
                    print(f"❌ Error processing {file_path}: {e}")
        
        self.parsed_files = processed_files
        
        if processed_files:
            print(f"✅ Successfully processed {len(processed_files)} files:")
            for file_name in processed_files:
                print(f"   • {file_name}")
        else:
            print("❌ No supported files found. Please add PDF, Excel, or Word files.")
        
        return self.extracted_data
    
    def generate_metrics_summary(self):
        """Generate a summary of all extracted metrics."""
        print("\n📊 EXTRACTED METRICS SUMMARY")
        print("=" * 60)
        
        if not self.extracted_data:
            print("❌ No data extracted. Please process files first.")
            return
        
        summary = {
            'aws_data_metrics': [],
            'geographic_data_metrics': [],
            'advertising_data_metrics': [],
            'prime_data_metrics': [],
            'operational_metrics_metrics': [],
            'total_files_processed': len(self.parsed_files)
        }
        
        # Aggregate metrics from all files
        for file_name, data in self.extracted_data.items():
            print(f"\n📄 {file_name}:")
            
            if 'metrics' in data:
                for metric_type, matches in data['metrics'].items():
                    if matches:  # If we found matches
                        print(f"  ✅ {metric_type}: {len(matches)} matches found")
                        summary[f"{metric_type}_metrics"].extend(matches)
                    else:
                        print(f"  ⚠️  {metric_type}: No matches found")
            
            if 'tables' in data:
                print(f"  📊 Tables extracted: {len(data['tables'])}")
                for table in data['tables']:
                    print(f"     • {table['type']} (Page {table.get('page', 'N/A')})")
        
        # Save summary to JSON
        summary_file = self.reports_dir / "parsed_data" / "metrics_summary.json"
        with open(summary_file, 'w') as f:
            json.dump(summary, f, indent=2, default=str)
        
        print(f"\n💾 Summary saved to: {summary_file}")
        return summary
    
    def export_to_dataframe(self):
        """Export extracted metrics to pandas DataFrame for analysis."""
        if not self.extracted_data:
            print("❌ No data to export. Please process files first.")
            return None
        
        # Create comprehensive DataFrame
        rows = []
        
        for file_name, data in self.extracted_data.items():
            if 'metrics' in data:
                for metric_type, matches in data['metrics'].items():
                    for match in matches:
                        rows.append({
                            'source_file': file_name,
                            'metric_type': metric_type,
                            'pattern': match.get('pattern', ''),
                            'matched_text': match.get('match', ''),
                            'extracted_value': match.get('value', ''),
                            'context': match.get('context', '')
                        })
        
        if rows:
            df = pd.DataFrame(rows)
            
            # Save to CSV
            csv_file = self.reports_dir / "parsed_data" / "extracted_metrics.csv"
            df.to_csv(csv_file, index=False)
            print(f"📊 Metrics exported to: {csv_file}")
            
            return df
        else:
            print("❌ No metrics data to export.")
            return None

def main():
    """Main function to demonstrate usage."""
    print("🚀 AMAZON INVESTOR REPORTS INTEGRATION SYSTEM")
    print("=" * 60)
    
    parser = AmazonInvestorReportsParser()
    
    # Check if files exist
    reports_dir = Path("investor_reports")
    if not reports_dir.exists() or not any(reports_dir.iterdir()):
        print("\n📂 SETUP INSTRUCTIONS:")
        print("=" * 40)
        print("1. Create folder: investor_reports/")
        print("2. Download Amazon investor files:")
        print("   • 10-K Annual Reports → investor_reports/10k_annual/")
        print("   • 10-Q Quarterly Reports → investor_reports/10q_quarterly/")
        print("   • Earnings Releases → investor_reports/earnings_releases/")
        print("   • Shareholder Letters → investor_reports/shareholder_letters/")
        print("\n📥 Download from: https://ir.aboutamazon.com/")
        print("   • Latest 10-K (Annual Report)")
        print("   • Latest 10-Q (Quarterly Report)")
        print("   • Recent Earnings Releases")
        print("   • CEO Shareholder Letters")
        print("\n🔄 Then run this script again to process the files.")
        return
    
    # Process files
    extracted_data = parser.process_all_files()
    
    if extracted_data:
        # Generate summary
        summary = parser.generate_metrics_summary()
        
        # Export to DataFrame
        df = parser.export_to_dataframe()
        
        print("\n✅ INTEGRATION COMPLETE!")
        print("Your investor report data is now integrated into the environment.")
        print("Use the exported CSV and JSON files in your econometric models.")
    else:
        print("\n❌ No files processed. Please add investor report files and try again.")

if __name__ == "__main__":
    main()
